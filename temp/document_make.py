import docx
from docx.oxml.ns import qn # 한글 폰트
from docx.enum.text import WD_ALIGN_PARAGRAPH # 중앙정렬
from docx.shared import Cm, Inches # 이미지 삽입시 길이 단위

import time # 파일생성시 이름에 날짜로 정렬

from docx2pdf import convert# docx를 pdf로 변환

# ++ from fpdf import FPDF (keep)
# ++ import pandas as pd (keep)

# 0. base
'''
* 사용시
from document_make import *

* 새로운 임포트 설치
pip3 install python-docx
pip3 install docx2pdf
pip3 install pypiwin32
pip3 install pywin32

* 아나콘다에서 설치 (pywin32 - 윈도우)
conda install -c anaconda pywin32
conda update -n base conda
conda update --all
python -m pip install --upgrade pip
'''

# 1. word (미사용)
'''
* 실행 예시 
document_make.word(1) # 1 == anything
'''

# 2. word_form
'''
* 전달 자료형
++ a

* 실행 예시 
document_make.word_form(1)
'''

# 3. convert_pdf
'''
* 실행 예시 
document_make.convert_pdf(1)

* 작동 원리
word_form에서 호출하는 식으로 작동
'''


class document_make:
    # 함수종류
    ### word : docx 생성 - 오직 코드로만 (미사용)
    ### word_form : 미리 정해진 양식을 사용하여 docx 생성
    ### convert_pdf : docx를 pdf로 변환

    # 전체적인 구조
    ### docx형식 파일을 불러와서 받은 데이터를 대입시킴
    ### docx를 pdf로 변환

    @staticmethod
    def word(insert_data):
        doc = docx.Document() # docx 생성

        para = doc.add_paragraph() 
        run = para.add_run('mini pot - title') # +본문
        run.font.size = docx.shared.Pt(30) # 폰트크기 
        run.bold = True # 볼트체 
        run.italic = True # 이텔릭체
        run.font.name = 'Cambria' # 폰트 설정
        # run._element.rPr.rFonts.set(qn('w:eastAsia'), '휴먼명조') # 한글 폰트 에러남
        last_paragraph = doc.paragraphs[-1]  # 이전 paragraph를 지정후 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER # 중앙정렬

        # 이미지 삽입
        ### ++ 이미지의 원래 비율에 맞개 줄이는 방법을 생각해야함
        ### ++ 아니면 규격을 정해놔서 거기에 맞춰야함
        doc.add_picture("test_image.png", width = Cm(13), height = Cm(8))

        # 단락 생성
        doc.add_paragraph('첫번째 단락', style='List Bullet')
        doc.add_paragraph('첫번째 순서 단락', style='List Number')

        # 공문서 : https://python-docx.readthedocs.io/en/latest/

        doc.save("time_plant.docx")

    @staticmethod
    def word_form(insert_data):
        print("START")

        # report_form을 가져와서 새로운 docx를 만들고 저장할 꺼임!
        doc = docx.Document("report_form.docx")

        # 변수에 전달 받은 값을 넣어서 자동화
        # ++ 변수 추가 설정 (받은 데이터, 내부 설계)
        title = "mini pot"
        user = "ZIMyMeMine"
        plant = "T-hub"

        # 제목 부분(para[0])에 title 변수를 대입 
        ## doc.paragraphs[0].text = title // 한줄로
        ###doc.add_heading("mini-pot", 0) // 단순 생성
        para = doc.paragraphs
        para[0].text = title
        
        
        #doc.add_paragraph("ttap", style = 'a') // 스타일이 a인 내용 ttap를 생성

        #para = doc.add_paragraph() // 해당 라인의 뒤에 추가로 생성
        #run = para.add_run('mini pot - title')

        # ++ 보여주는 타입을 2개로 나누어서 일반:그래프, 고급:표



        # 디버깅 : word 문서를 번호 : 내용으로 확인
        for x, paragraph in enumerate(doc.paragraphs):
            print(str(x) + " : " + paragraph.text)

        # 생성 날짜와 유저,식물 정보로 사진을 저장
        ## ex) ZIMyMeMine_20210701-091731_T-hub
        ## ++ 지금은 하위 폴더(report_result)에 docx저장하고 pdf 변환 추후 경로 수정 필요
        time_file = time.strftime('%Y%m%d-%H%M%S', time.localtime(time.time()))
        file_path = ".\\report_result\\"
        file_name = user + "_" + time_file + "_" + plant
        doc.save(file_path + file_name + ".docx")

        # 생성된 파일이름을 pdf로 변환하기 위해 전달
        #document_make.convert_pdf(file_name) ++

    @staticmethod
    def convert_pdf(convert_file_name):
        
        # 받은 파일이름에 확장자랑 경로를 붙여줌
        docx = convert_file_name + ".docx"
        pdf = convert_file_name + ".pdf"
        docx_path = ".\\report_result\\" # !! ""안에 있으니까 \\ 이렇게 두번 해줘야함 
        pdf_path = ".\\report_result\\" # !! docx랑 pdf랑 동일 경로에 설정하면 안됨(윈도우) 리눅스는 아직 실험X but 하위 폴더일때는 가능

        # 변환
        convert(docx_path+docx, pdf_path+pdf)

        # 이부분에서 파일 경로를 전달할 send로 보내주면 됨
        ## ++ 상대경로를 줄지 아니면 파일 이름만 줘서 경로는 send에서 고정으로 할지는 send 생성후 다시
        print(pdf_path + pdf)

        
# test space ----------------------------------------
''''''
document_make.word_form(1)
